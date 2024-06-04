import tkinter as tk
from tkinter import filedialog
from tkinter import scrolledtext
from docx import Document
from collections import Counter
import re
from abc import ABC, abstractmethod

class Logger:
    _instance = None

    def __new__(cls, *args, **kwargs):
        if not cls._instance:
            cls._instance = super(Logger, cls).__new__(cls, *args, **kwargs)
            cls._instance.logs = []
        return cls._instance

    def log(self, message):
        self.logs.append(message)
        print(message)  

    def get_logs(self):
        return self.logs

logger = Logger()

class TextAnalysisStrategy(ABC):
    @abstractmethod
    def analyze(self, text):
        pass

class WordFrequencyStrategy(TextAnalysisStrategy):
    def __init__(self, word):
        self.word = word.lower()

    def analyze(self, text):
        words = re.findall(r'\b\w+\b', text.lower())
        total_words = len(words)
        word_counts = Counter(words)
        frequency = word_counts[self.word] / total_words if total_words > 0 else 0
        result = f"Відносна частота вживання '{self.word}': {frequency:.4f}"
        logger.log(result)
        return result

class SentenceCountStrategy(TextAnalysisStrategy):
    def analyze(self, text):
        sentences = re.split(r'[.!?]+', text)
        total_sentences = len([s for s in sentences if s.strip()])
        result = f"Кількість речень: {total_sentences}"
        logger.log(result)
        return result

class PunctuationCountStrategy(TextAnalysisStrategy):
    def analyze(self, text):
        punctuation_marks = re.findall(r'[.,!?;:—]', text)
        total_punctuation = len(punctuation_marks)
        result = f"Кількість розділових знаків: {total_punctuation}"
        logger.log(result)
        return result

def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc = Document(file_path)
        text_editor.delete(1.0, tk.END)  
        for paragraph in doc.paragraphs:
            text_editor.insert(tk.INSERT, paragraph.text + '\n')  
        logger.log(f"Відкрито файл: {file_path}")

def analyze_text(): # стратегія вибирається на основі користувацького введення і застосовується до тексту 
    text = text_editor.get(1.0, tk.END).strip()
    strategy_name = strategy_var.get()
    if strategy_name == "Відносна частота":
        strategy = WordFrequencyStrategy(word_entry.get())
    elif strategy_name == "Кількість речень":
        strategy = SentenceCountStrategy()
    elif strategy_name == "Кількість розділових знаків":
        strategy = PunctuationCountStrategy()
    else:
        result_label.config(text="Виберіть стратегію аналізу")
        return

    result = strategy.analyze(text)
    result_label.config(text=result)

def save_file():
    text = text_editor.get(1.0, tk.END).strip()
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
        logger.log(f"Збережено файл: {file_path}")

root = tk.Tk()
root.title("Текстовий Редактор")

open_button = tk.Button(root, text="Відкрити файл", command=open_file)
open_button.pack(pady=10)

text_editor = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=80, height=20)
text_editor.pack(pady=10, padx=10)

word_entry = tk.Entry(root)
word_entry.pack(pady=5)

strategy_var = tk.StringVar(value="Відносна частота")
strategy_menu = tk.OptionMenu(root, strategy_var, "Відносна частота", "Кількість речень", "Кількість розділових знаків")
strategy_menu.pack(pady=5)

analyze_button = tk.Button(root, text="Провести аналіз", command=analyze_text)
analyze_button.pack(pady=5)

save_button = tk.Button(root, text="Зберегти файл", command=save_file)
save_button.pack(pady=5)

result_label = tk.Label(root, text="")
result_label.pack(pady=5)

root.mainloop()



